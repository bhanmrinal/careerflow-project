"""
Resume Parser Service.

Handles parsing of PDF and DOCX resume files, extracting text and
identifying sections using LLM-based analysis.
"""

import io
import re
from pathlib import Path
from typing import BinaryIO
from uuid import uuid4

from backend.app.core.config import get_settings
from backend.app.core.llm import get_llm
from backend.app.models.resume import Resume, ResumeSection, SectionType


class ResumeParserService:
    """Service for parsing resume files and extracting structured content."""

    SECTION_PATTERNS = {
        SectionType.CONTACT: r"(contact|personal\s*info|email|phone)",
        SectionType.SUMMARY: r"(summary|objective|profile|about)",
        SectionType.EXPERIENCE: r"(experience|employment|work\s*history|professional)",
        SectionType.EDUCATION: r"(education|academic|qualification|degree)",
        SectionType.SKILLS: r"(skills|technical|competencies|expertise)",
        SectionType.PROJECTS: r"(projects|portfolio|work\s*samples)",
        SectionType.CERTIFICATIONS: r"(certification|license|credential)",
        SectionType.LANGUAGES: r"(languages|linguistic)",
    }

    def __init__(self):
        self.settings = get_settings()

    async def parse_file(self, file: BinaryIO, filename: str, user_id: str) -> Resume:
        """
        Parse a resume file and extract structured content.

        Args:
            file: File-like object containing the resume.
            filename: Original filename.
            user_id: ID of the user uploading the resume.

        Returns:
            Parsed Resume object.

        Raises:
            ValueError: If file type is not supported.
        """
        extension = Path(filename).suffix.lower().lstrip(".")

        if extension not in self.settings.allowed_extensions:
            raise ValueError(
                f"Unsupported file type: {extension}. "
                f"Allowed: {', '.join(self.settings.allowed_extensions)}"
            )

        if extension == "pdf":
            raw_text = self._parse_pdf(file)
        elif extension == "docx":
            raw_text = self._parse_docx(file)
        else:
            raise ValueError(f"Unsupported file type: {extension}")

        sections = await self._extract_sections(raw_text)

        resume = Resume(
            id=str(uuid4()),
            user_id=user_id,
            filename=filename,
            raw_text=raw_text,
            sections=sections,
            metadata={
                "original_extension": extension,
                "character_count": len(raw_text),
                "sections_count": len(sections),
            },
        )

        return resume

    def _parse_pdf(self, file: BinaryIO) -> str:
        """Extract text from a PDF file."""
        import pdfplumber

        text_parts = []
        file_bytes = file.read()

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        return "\n\n".join(text_parts)

    def _parse_docx(self, file: BinaryIO) -> str:
        """Extract text from a DOCX file."""
        from docx import Document

        file_bytes = file.read()
        doc = Document(io.BytesIO(file_bytes))

        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts)

    async def _extract_sections(self, raw_text: str) -> list[ResumeSection]:
        """
        Extract sections from raw resume text using LLM.

        Args:
            raw_text: Raw text extracted from resume.

        Returns:
            List of identified resume sections.
        """
        llm = get_llm(temperature=0.1)

        prompt = f"""Analyze the following resume text and identify distinct sections.
For each section, provide:
1. The section type (one of: contact, summary, experience, education, skills, projects, certifications, languages, other)
2. The section title as it appears in the resume
3. The content of that section

Resume Text:
{raw_text}

Respond in the following format for each section:
SECTION_TYPE: <type>
TITLE: <title>
CONTENT:
<content>
---

Be thorough and capture all sections present in the resume."""

        response = await llm.ainvoke(prompt)
        sections = self._parse_section_response(response.content)

        if not sections:
            sections = self._fallback_section_extraction(raw_text)

        return sections

    def _parse_section_response(self, response: str) -> list[ResumeSection]:
        """Parse LLM response into ResumeSection objects."""
        sections = []
        section_blocks = response.split("---")
        order = 0

        for block in section_blocks:
            block = block.strip()
            if not block:
                continue

            section_type_match = re.search(
                r"SECTION_TYPE:\s*(\w+)", block, re.IGNORECASE
            )
            title_match = re.search(
                r"TITLE:\s*(.+?)(?=\n|CONTENT:)", block, re.IGNORECASE
            )
            content_match = re.search(
                r"CONTENT:\s*(.+)", block, re.IGNORECASE | re.DOTALL
            )

            if section_type_match and content_match:
                type_str = section_type_match.group(1).lower()
                section_type = self._map_section_type(type_str)

                title = (
                    title_match.group(1).strip() if title_match else type_str.title()
                )
                content = content_match.group(1).strip()

                sections.append(
                    ResumeSection(
                        section_type=section_type,
                        title=title,
                        content=content,
                        order=order,
                    )
                )
                order += 1

        return sections

    def _map_section_type(self, type_str: str) -> SectionType:
        """Map a string to SectionType enum."""
        type_mapping = {
            "contact": SectionType.CONTACT,
            "summary": SectionType.SUMMARY,
            "objective": SectionType.SUMMARY,
            "profile": SectionType.SUMMARY,
            "experience": SectionType.EXPERIENCE,
            "work": SectionType.EXPERIENCE,
            "employment": SectionType.EXPERIENCE,
            "education": SectionType.EDUCATION,
            "skills": SectionType.SKILLS,
            "technical": SectionType.SKILLS,
            "projects": SectionType.PROJECTS,
            "certifications": SectionType.CERTIFICATIONS,
            "certificates": SectionType.CERTIFICATIONS,
            "languages": SectionType.LANGUAGES,
        }
        return type_mapping.get(type_str.lower(), SectionType.OTHER)

    def _fallback_section_extraction(self, raw_text: str) -> list[ResumeSection]:
        """
        Fallback section extraction using regex patterns.

        Used when LLM extraction fails.
        """
        sections = []
        lines = raw_text.split("\n")
        current_section = None
        current_content = []
        order = 0

        for line in lines:
            line_lower = line.lower().strip()

            detected_type = None
            for section_type, pattern in self.SECTION_PATTERNS.items():
                if re.search(pattern, line_lower):
                    detected_type = section_type
                    break

            if detected_type:
                if current_section and current_content:
                    sections.append(
                        ResumeSection(
                            section_type=current_section,
                            title=current_content[0]
                            if current_content
                            else str(current_section.value),
                            content="\n".join(current_content[1:])
                            if len(current_content) > 1
                            else "",
                            order=order,
                        )
                    )
                    order += 1

                current_section = detected_type
                current_content = [line.strip()]
            elif current_section:
                current_content.append(line)

        if current_section and current_content:
            sections.append(
                ResumeSection(
                    section_type=current_section,
                    title=current_content[0]
                    if current_content
                    else str(current_section.value),
                    content="\n".join(current_content[1:])
                    if len(current_content) > 1
                    else "",
                    order=order,
                )
            )

        if not sections:
            sections.append(
                ResumeSection(
                    section_type=SectionType.OTHER,
                    title="Full Resume",
                    content=raw_text,
                    order=0,
                )
            )

        return sections
